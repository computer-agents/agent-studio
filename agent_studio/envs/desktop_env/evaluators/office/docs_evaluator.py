import logging
import re
from io import BytesIO
from typing import Any

from docx import Document, document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_TAB_ALIGNMENT
from docx.shared import RGBColor
from odf.opendocument import load
from odf.text import P, Span
from PIL import Image
from rapidfuzz import fuzz
from skimage.color import deltaE_ciede2000, rgb2lab

from agent_studio.config import Config
from agent_studio.envs.desktop_env.evaluators.evaluator import (
    Evaluator,
    FeedbackException,
    evaluation_handler,
)

config = Config()
logger = logging.getLogger(__name__)


def _compare_docx_files(file1, file2, options: dict = {}):
    ignore_blanks = options.get("ignore_blanks", True)
    ignore_case = options.get("ignore_case", False)
    ignore_order = options.get("ignore_order", False)
    content_only = options.get("content_only", False)
    similarity_threshold = options.get("similarity_threshold", 1.0)


    def get_paragraph_texts_odt(document):
        paragraphs = document.getElementsByType(P)
        paragraph_texts = []
        for paragraph in paragraphs:
            text_parts = []
            for node in paragraph.childNodes:
                if node.nodeType == node.TEXT_NODE:
                    text_parts.append(node.data)
                elif node.nodeType == node.ELEMENT_NODE and node.tagName == "text:span":
                    # Assuming direct text content in <text:span>, for simplicity
                    for child in node.childNodes:
                        if child.nodeType == child.TEXT_NODE:
                            text_parts.append(child.data)
            paragraph_texts.append("".join(text_parts))
        return paragraph_texts

    # Determine file types and load documents
    if file1.endswith(".docx") and file2.endswith(".docx"):
        try:
            doc1 = Document(file1)
            doc2 = Document(file2)
        except Exception as e:
            logger.error(f"Error: {e}")
            return 0
        doc1_paragraphs = [p.text for p in doc1.paragraphs]
        doc2_paragraphs = [p.text for p in doc2.paragraphs]
        if ignore_order:
            doc1_paragraphs = sorted(doc1_paragraphs)
            doc2_paragraphs = sorted(doc2_paragraphs)
    elif file1.endswith(".odt") and file2.endswith(".odt"):
        try:
            doc1 = load(file1)
            doc2 = load(file2)
        except Exception as e:
            logger.error(f"Error: {e}")
            raise FeedbackException(f"Error loading documents: {e}")
        doc1_paragraphs = get_paragraph_texts_odt(doc1)
        doc2_paragraphs = get_paragraph_texts_odt(doc2)
        if ignore_order:
            doc1_paragraphs = sorted(doc1_paragraphs)
            doc2_paragraphs = sorted(doc2_paragraphs)
    else:
        # Unsupported file types or mismatch
        raise FeedbackException("Unsupported file types or mismatch between file types.")

    if content_only:
        # Compare the content of the documents
        text1 = re.sub(r"\s+", " ", "\n".join(doc1_paragraphs)).strip()
        text2 = re.sub(r"\s+", " ", "\n".join(doc2_paragraphs)).strip()
        if ignore_case:
            text1, text2 = text1.lower(), text2.lower()
        similarity = fuzz.ratio(text1, text2) / 100.0
        if similarity < similarity_threshold:
            raise FeedbackException(f"Documents are different, Text similarity is {similarity:.2f}, {text1} != {text2}")

    # Process and compare documents
    if ignore_blanks:
        text1 = re.sub(r"\s+", " ", "\n".join(doc1_paragraphs)).strip()
        text2 = re.sub(r"\s+", " ", "\n".join(doc2_paragraphs)).strip()
        if ignore_case:
            text1, text2 = text1.lower(), text2.lower()
        if text1 != text2:
            raise FeedbackException(f"Documents are different, {text1} != {text2}")
    else:
        if len(doc1_paragraphs) != len(doc2_paragraphs):
            raise FeedbackException(f"Number of paragraphs is different, {len(doc1_paragraphs)} != {len(doc2_paragraphs)}")
        # Compare each paragraph
        for p1, p2 in zip(doc1_paragraphs, doc2_paragraphs):
            if ignore_case:
                p1, p2 = p1.lower(), p2.lower()
            if p1 != p2:
                raise FeedbackException(f"Documents are different, {p1} != {p2}")


class DocsCalcEvaluator(Evaluator):
    name: str = "docs"

    @evaluation_handler("compare_line_spacing")
    def compare_line_spacing(self, docx_file1, docx_file2):
        _compare_docx_files(docx_file1, docx_file2)

        try:
            doc1 = Document(docx_file1)
            doc2 = Document(docx_file2)
        except Exception as e:
            logger.error(f"Error: {e}")
            raise FeedbackException("Error loading documents")

        if len(doc1.paragraphs) != len(doc2.paragraphs):
            raise FeedbackException("Number of paragraphs is different")

        # Compare each paragraph line spacing
        for para1, para2 in zip(doc1.paragraphs, doc2.paragraphs):
            spacing1 = para1.paragraph_format.line_spacing
            spacing2 = para2.paragraph_format.line_spacing

            if spacing1 != spacing2:
                raise FeedbackException(
                    f"Line spacing of paragraph '{para1.text}' is different"
                )

    @evaluation_handler("check_tabstops")
    def check_tabstops(self, docx_ref, docx_file, kwargs):
        try:
            doc1: document.Document = Document(docx_ref)
            doc2: document.Document = Document(docx_file)
        except Exception as e:
            logger.error(f"Error: {e}")
            raise FeedbackException(f"Error loading documents: {e}")

        para1 = [p for p in doc1.paragraphs if p.text.strip()]
        para2 = [p for p in doc2.paragraphs if p.text.strip()]
        if len(para1) != len(para2):
            raise FeedbackException(
                f"Number of paragraphs is different, {len(para1)} != {len(para2)}")

        if kwargs.get("word_number_split_by_tabstop", None) is not None:
            number = kwargs["word_number_split_by_tabstop"]
            index = kwargs.get("index", 0)
            for p2 in para2:
                splits = p2.text.split("\t")
                if len(splits) == 0:
                    raise FeedbackException(f"Paragraph {p2.text} is empty")
                words = list(
                    filter(lambda x: x.strip(), re.split(r"\s", splits[index]))
                )
                if len(words) != number:
                    raise FeedbackException(
                        f"Words: [{splits[index]}] has {len(words)} words, not {number}")

        section = doc2.sections[0]
        if section.page_width is None:
            raise FeedbackException("Page width is None")
        if section.left_margin is None:
            raise FeedbackException("Left margin is None")
        if section.right_margin is None:
            raise FeedbackException("Right margin is None")
        paragraph_width = (
            section.page_width - section.left_margin - section.right_margin
        )

        def ignore_tabs(x): return x.alignment == WD_TAB_ALIGNMENT.CLEAR or (
            x.alignment == WD_TAB_ALIGNMENT.LEFT and x.position == 0
        )
        minus = 0.0
        for p1, p2 in zip(para1, para2):
            # filter CLEAR tabstop and default left-0 tabstop
            tabs1 = [
                tst for tst in p1.paragraph_format.tab_stops if not ignore_tabs(tst)
            ]
            tabs2 = [
                tst for tst in p2.paragraph_format.tab_stops if not ignore_tabs(tst)
            ]
            if len(tabs1) != len(tabs2):
                return 0.0
            difference = 0.0
            for t1, t2 in zip(tabs1, tabs2):
                if t1.alignment != t2.alignment:
                    return 0.0
                difference += abs(t1.position - t2.position)
            minus += difference / paragraph_width
        score = 1 - (minus / len(para1))
        if not score:
            raise FeedbackException("Tabstops are different")

    @evaluation_handler("compare_subscript_contains")
    def compare_subscript_contains(self, docx_file1, docx_file2):
        try:
            doc1 = Document(docx_file1)
            doc2 = Document(docx_file2)
        except Exception as e:
            logger.error(f"Error: {e}")
            raise FeedbackException("Error loading documents")

        for para1, para2 in zip(doc1.paragraphs, doc2.paragraphs):
            for run1, run2 in zip(para1.runs, para2.runs):
                # check if two paras both contain subscript
                if run1.font.subscript and run2.font.subscript:
                    return

        raise FeedbackException("Subscript is missing")

    @evaluation_handler("has_page_numbers_in_footers")
    def has_page_numbers_in_footers(self, docx_file):
        try:
            doc = Document(docx_file)
        except Exception as e:
            logger.error(f"Error: {e}")
            raise FeedbackException("Error loading document")

        for section in doc.sections:
            footer = section.footer
            if footer is None:
                raise FeedbackException("Footer is missing")
            footer_text = footer.paragraphs[0].text if footer.paragraphs else ""
            if not any(char.isdigit() for char in footer_text):
                # if no digit in footer, then no page number
                raise FeedbackException("Page number is missing in footer")

    @evaluation_handler("compare_font_names")
    def compare_font_names(self, docx_file, rules: dict[str, Any]):
        try:
            doc = Document(docx_file)
        except Exception as e:
            logger.error(f"Error: {e}")
            raise FeedbackException("Error loading document")

        expected_font = rules["font_name"]

        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                font_name = run.font.name
                if font_name != expected_font:
                    raise FeedbackException(f"Font name is not {expected_font}")

    @evaluation_handler("is_first_line_centered")
    def is_first_line_centered(self, docx_file):
        try:
            doc = Document(docx_file)
        except Exception as e:
            logger.error(f"Error: {e}")
            raise FeedbackException("Error loading document")

        first_paragraph = doc.paragraphs[0]

        # check if the first line is center justified
        if first_paragraph.paragraph_format.alignment != WD_PARAGRAPH_ALIGNMENT.CENTER:
            raise FeedbackException("First line is not center justified")

    @evaluation_handler("compare_insert_equation")
    def compare_insert_equation(self, docx_file1, docx_file2):
        _compare_docx_files(docx_file1, docx_file2)

        try:
            doc1 = Document(docx_file1)
            doc2 = Document(docx_file2)
        except Exception as e:
            logger.error(f"Error: {e}")
            raise FeedbackException("Error loading documents")

        # Compare each paragraph if it contains equation
        equ_list1 = []
        equ_list2 = []
        for para1 in doc1.paragraphs:
            para_xml = para1._element.xml
            if "<m:oMath" in para_xml:
                equ_list1.append(para_xml)

        for para2 in doc2.paragraphs:
            para_xml = para2._element.xml
            if "<m:oMath" in para_xml:
                equ_list2.append(para_xml)

        if not equ_list1 or not equ_list2:
            raise FeedbackException("Equation is missing")

        # Compare equations
        # TODO: compare the content of the equations
        if len(equ_list1) != len(equ_list2):
            raise FeedbackException(f"Number of equations is different, {len(equ_list1)} != {len(equ_list2)}")
        # for equ1, equ2 in zip(equ_list1, equ_list2):
        #     if equ1 != equ2:
        #         raise FeedbackException(f"Equations are different, {equ1} != {equ2}")

    @evaluation_handler("compare_docx_tables")
    def compare_docx_tables(self, docx_file1, docx_file2):
        try:
            doc1 = Document(docx_file1)
        except Exception as e:
            logger.error(f"Error: {e}")
            raise FeedbackException(f"Error loading documents: {e}")
        try:
            doc2 = Document(docx_file2)
        except Exception as e:
            logger.error(f"Error: {e}")
            raise FeedbackException(f"Error loading documents: {e}")

        # get list of tables in docx
        tables1 = doc1.tables
        tables2 = doc2.tables

        if len(tables1) != len(tables2):
            raise FeedbackException("Number of tables is different")

        # Compare each table content
        for table1, table2 in zip(tables1, tables2):
            if len(table1.rows) != len(table2.rows) or len(table1.columns) != len(
                table2.columns
            ):
                raise FeedbackException("Table structure is different")

            # Compare each cell
            for i in range(len(table1.rows)):
                for j in range(len(table1.columns)):
                    if table1.cell(i, j).text.strip() != table2.cell(i, j).text.strip():
                        raise FeedbackException("Table content is different")

    @evaluation_handler("check_highlighted_words")
    def check_highlighted_words(self, file_path1, file_path2):
        _compare_docx_files(file_path1, file_path2)

        doc = load(file_path1)
        highlighted = False

        for span in doc.getElementsByType(Span):
            style_name = span.getAttribute("stylename")
            if style_name:
                for automatic_style in doc.automaticstyles.childNodes:
                    if automatic_style.getAttribute("name") == style_name:
                        for property in automatic_style.childNodes:
                            if property.getAttribute("backgroundcolor") == "#ffff00":
                                highlighted = True
                                break
                if highlighted:
                    break

        if highlighted:
            raise FeedbackException("Highlighted words are present")

    @evaluation_handler("compare_contains_image")
    def compare_contains_image(self, docx_file1, docx_file2):
        try:
            doc1 = Document(docx_file1)
            doc2 = Document(docx_file2)
        except Exception as e:
            logger.error(f"Error: {e}")
            raise FeedbackException("Error loading documents")

        for para1, para2 in zip(doc1.paragraphs, doc2.paragraphs):
            for run1, run2 in zip(para1.runs, para2.runs):
                if ("graphicData" in run1._element.xml) != ("graphicData" in run2._element.xml):
                    raise FeedbackException(
                        f"Image is missing, {run1._element.xml} != {run2._element.xml}")

    @evaluation_handler("compare_docx_lines")
    def compare_docx_lines(self, file1, file2):
        # Read the text of the document, line by line
        try:
            doc1 = Document(file1)
            doc2 = Document(file2)
        except Exception as e:
            logger.error(f"Error: {e}")
            raise FeedbackException("Error loading documents")

        doc1_lines = [p.text.strip() for p in doc1.paragraphs if p.text.strip()]
        doc2_lines = [p.text.strip() for p in doc2.paragraphs if p.text.strip()]

        # Convert the list of lines to sets and compare
        if set(doc1_lines) != set(doc2_lines):
            raise FeedbackException("Documents are different")

    @evaluation_handler("evaluate_strike_through_last_paragraph")
    def evaluate_strike_through_last_paragraph(self, docx_ref, docx_file):
        _compare_docx_files(docx_ref, docx_file)

        try:
            document = Document(docx_ref)
        except Exception as e:
            logger.error(f"Error: {e}")
            raise FeedbackException("Error loading document")

        # Get the last paragraph
        last_paragraph = document.paragraphs[-1]

        # Check if any run in the last paragraph has strike-through formatting
        for run in last_paragraph.runs:
            if not run.font.strike:
                raise FeedbackException(
                    "Strike-through formatting is missing"
                )  # At least one word does not have strike-through formatting

        # All words in the last paragraph have strike-through formatting

    @evaluation_handler("evaluate_colored_words_in_tables")
    def evaluate_colored_words_in_tables(self, docx_ref, docx_file):
        _compare_docx_files(docx_ref, docx_file)

        try:
            document = Document(docx_file)
        except Exception as e:
            logger.error(f"Error: {e}")
            raise FeedbackException("Error loading document")

        def _calculate_color_difference(rgb1, rgb2):
            srgb1 = [rgb1[0] / 255.0, rgb1[1] / 255.0, rgb1[2] / 255.0]
            srgb2 = [rgb2[0] / 255.0, rgb2[1] / 255.0, rgb2[2] / 255.0]
            return srgb1.index(max(srgb1)) != srgb2.index(max(srgb2))

        for table in document.tables:
            # Iterate through rows and cells in the table
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            word = run.text
                            if word:
                                first_letter = word[0].lower()

                                if run.font.color.rgb is None or (
                                    first_letter in "aeiou"
                                    and _calculate_color_difference(
                                        run.font.color.rgb, [255, 0, 0]
                                    )
                                ):
                                    raise FeedbackException(
                                        f"Vowel-colored words should be red, {run.font.color.rgb} != {[255, 0, 0]}, Word: {word}"
                                    )
                                elif run.font.color.rgb is None or (
                                    first_letter not in "aeiou"
                                    and _calculate_color_difference(
                                        run.font.color.rgb, [0, 0, 255]
                                    )
                                ):
                                    raise FeedbackException(
                                        f"Non-vowel-colored words should be blue, {run.font.color.rgb} != {[0, 0, 255]}, Word: {word}"
                                    )

    @evaluation_handler("compare_docx_files")
    def compare_docx_files(self, docx_file1, docx_file2, options={}):
        _compare_docx_files(docx_file1, docx_file2, options)

    @evaluation_handler("compare_docx_images")
    def compare_docx_images(self, docx_file1, docx_file2):
        try:
            doc1 = Document(docx_file1)
        except Exception as e:
            logger.error(f"Error: {e}")
            raise FeedbackException(f"Error loading documents: {e}")
        try:
            doc2 = Document(docx_file2)
        except Exception as e:
            logger.error(f"Error: {e}")
            raise FeedbackException(f"Error loading documents: {e}")

        def extract_images(doc):
            images = []
            for rel in doc.part.rels.values():
                if "image" in rel.reltype:
                    img_data = rel.target_part.blob
                    images.append(BytesIO(img_data))
            return images

        images1 = extract_images(doc1)
        images2 = extract_images(doc2)
        if len(images1) != len(images2):
            raise FeedbackException("Number of images is different")
        for img1, img2 in zip(images1, images2):
            if Image.open(img1).tobytes() != Image.open(img2).tobytes():
                raise FeedbackException("Images are different")

    @evaluation_handler("compare_references")
    def compare_references(self, docx_file1, docx_file2, options={}):
        reference_indicator = options.get("reference_indicator", "References")
        reference_base_result = options.get("reference_base_result", 0.5)

        # Determine file types and load documents
        if docx_file1.endswith(".docx") and docx_file2.endswith(".docx"):
            try:
                doc1 = Document(docx_file1)
                doc2 = Document(docx_file2)
            except Exception as e:
                logger.error(f"Error: {e}")
                raise FeedbackException("Error loading documents")

            doc1_paragraphs = [p.text for p in doc1.paragraphs]
            doc2_paragraphs = [p.text for p in doc2.paragraphs]
        else:
            # Unsupported file types or mismatch
            print("Unsupported file types or mismatch between file types.")
            raise FeedbackException(
                "Unsupported file types or mismatch between file types."
            )

        # Find the references section in the paragraphs, find the idx of the last reference_indicator in the paragraph list  # noqa: E501
        ref1_idx = (
            doc1_paragraphs.index(reference_indicator)
            if reference_indicator in doc1_paragraphs
            else -1
        )
        ref2_idx = (
            doc2_paragraphs.index(reference_indicator)
            if reference_indicator in doc2_paragraphs
            else -1
        )

        if ref1_idx == -1 or ref2_idx == -1:
            raise FeedbackException("References section is missing")

        # split the reference section into reference items, and remove the empty string items  # noqa: E501
        ref1 = [p for p in doc1_paragraphs[ref1_idx + 1:] if p.strip()]
        ref2 = [p for p in doc2_paragraphs[ref2_idx + 1:] if p.strip()]

        # Compare the references

        if len(ref1) != len(ref2):
            raise FeedbackException("Number of references is different")

        total_similarity = 0
        for r1, r2 in zip(ref1, ref2):
            # fuzzy match the references
            similarity = fuzz.ratio(r1, r2) / 100.0
            total_similarity += similarity

        result = total_similarity / len(ref1)
        if result < reference_base_result:
            raise FeedbackException(f"References are different, {result} < {reference_base_result}")
            # return (result - reference_base_result) / (1 - reference_base_result)
